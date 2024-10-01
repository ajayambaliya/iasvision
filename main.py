import requests
from bs4 import BeautifulSoup
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import pypandoc
from deep_translator import GoogleTranslator  # Import the deep translator

# Telegram bot credentials
BOT_TOKEN = "1637529837:AAFraGS_WwfTV8rj9XOhBy7PoxnbnVXBVEM"
CHAT_ID = "@gujtest"  # Use the channel handle, like "@yourchannel" or the actual chat ID

# Define base URL
base_url = "https://visionias.in/current-affairs/"

# Get the previous date in 'yyyy-mm-dd' format
previous_date = (datetime.now() - timedelta(days=1)).strftime('%Y-%m-%d')

# Send GET request to the base URL
response = requests.get(base_url)
soup = BeautifulSoup(response.content, 'html.parser')

# Find all anchor tags with href containing the previous date
anchors = soup.find_all('a', href=lambda href: href and previous_date in href)

# Create a set to store unique URLs (set ensures uniqueness)
unique_urls = set()

for a in anchors:
    href = a['href']
    if href.startswith('/current-affairs/'):
        # Remove '/current-affairs/' from the href to get the correct URL format
        correct_href = href.replace('/current-affairs/', '', 1)
        full_url = base_url.rstrip('/') + '/' + correct_href
        unique_urls.add(full_url)  # Adding to set ensures no duplicates

# Initialize a Word Document
doc = Document()

# Create an instance of the GoogleTranslator to translate to Gujarati
translator = GoogleTranslator(source='auto', target='gu')

# Dictionary to store the original content for later writing
original_content = {}

# Function to scrape content from the provided URL and translate to Gujarati
def scrap_and_create_doc(url):
    print(f"Scraping: {url}")
    page_response = requests.get(url)
    page_soup = BeautifulSoup(page_response.content, 'html.parser')

    # Find the specific content area div
    content_area = page_soup.find('div', class_="flex flex-col w-full mt-6 lg:mt-0")

    if content_area:
        # Extract the title from the <h1> tag
        title = content_area.find('h1').get_text()
        translated_title = translator.translate(title)  # Translate the title to Gujarati
        # Add title to the Word document (Gujarati content)
        doc.add_heading(translated_title, 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Store the original title in the dictionary for later use
        original_content[url] = {"title": title, "content": []}

        # Find article content under the <div id="article-content">
        article_content = content_area.find('div', id='article-content')

        if article_content:
            # Loop through each child element in the article content
            for element in article_content.find_all(recursive=False):
                if element.name == 'p':  # For paragraphs
                    paragraph_text = element.get_text()
                    translated_paragraph = translator.translate(paragraph_text)  # Translate paragraph
                    # Add Gujarati paragraph to the Word document
                    paragraph = doc.add_paragraph(translated_paragraph)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    # Store original paragraph for later use
                    original_content[url]["content"].append({"type": "p", "text": paragraph_text})

                elif element.name == 'h2':  # For sub-headings
                    sub_heading_text = element.get_text()
                    translated_sub_heading = translator.translate(sub_heading_text)  # Translate sub-heading
                    # Add Gujarati sub-heading to the Word document
                    doc.add_heading(translated_sub_heading, level=2)
                    # Store original sub-heading for later use
                    original_content[url]["content"].append({"type": "h2", "text": sub_heading_text})

                elif element.name == 'ul':  # For lists
                    list_items = []
                    for li in element.find_all('li'):
                        list_item_text = li.get_text()
                        translated_list_item = translator.translate(list_item_text)  # Translate list item
                        # Add Gujarati list item to the Word document
                        doc.add_paragraph(translated_list_item, style='List Bullet')
                        list_items.append(list_item_text)
                    # Store original list items for later use
                    original_content[url]["content"].append({"type": "ul", "items": list_items})

                elif element.name == 'ol':  # For ordered lists
                    list_items = []
                    for li in element.find_all('li'):
                        list_item_text = li.get_text()
                        translated_list_item = translator.translate(list_item_text)  # Translate list item
                        # Add Gujarati ordered list item to the Word document
                        doc.add_paragraph(translated_list_item, style='List Number')
                        list_items.append(list_item_text)
                    # Store original ordered list items for later use
                    original_content[url]["content"].append({"type": "ol", "items": list_items})

# Iterate over unique URLs and scrape content
for url in unique_urls:
    scrap_and_create_doc(url)

# After writing all Gujarati content, add a heading to separate the original content
doc.add_page_break()  # Add a page break before original content
doc.add_heading('Original Content in English', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Write original content to the document
for url, content_data in original_content.items():
    # Add the original title
    doc.add_heading(content_data["title"], level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the original content (paragraphs, headings, lists, etc.)
    for content_item in content_data["content"]:
        if content_item["type"] == "p":
            paragraph = doc.add_paragraph(content_item["text"])
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif content_item["type"] == "h2":
            doc.add_heading(content_item["text"], level=2)
        elif content_item["type"] == "ul":
            for item in content_item["items"]:
                doc.add_paragraph(item, style='List Bullet')
        elif content_item["type"] == "ol":
            for item in content_item["items"]:
                doc.add_paragraph(item, style='List Number')

# Save the document
file_name = f"visionias_current_affairs_{previous_date}.docx"
doc.save(file_name)
print(f"Document saved as {file_name}")

# Convert DOCX to PDF using pypandoc
def convert_docx_to_pdf(docx_file, pdf_file):
    output = pypandoc.convert_file(docx_file, 'pdf', outputfile=pdf_file)
    assert output == "", "Conversion failed"
    print(f"Document converted to PDF: {pdf_file}")

# Convert DOCX to PDF
pdf_file_name = f"visionias_current_affairs_{previous_date}.pdf"
convert_docx_to_pdf(file_name, pdf_file_name)

# Send the PDF file to Telegram
def send_file_to_telegram(pdf_file):
    url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendDocument"
    with open(pdf_file, 'rb') as file:
        response = requests.post(url, data={'chat_id': CHAT_ID}, files={'document': file})
    
    if response.status_code == 200:
        print(f"PDF file sent to Telegram channel successfully.")
    else:
        print(f"Failed to send PDF file. Status code: {response.status_code}, Response: {response.text}")

# Send the generated PDF to Telegram
send_file_to_telegram(pdf_file_name)

# Optionally, cleanup the local files
os.remove(file_name)
os.remove(pdf_file_name)
